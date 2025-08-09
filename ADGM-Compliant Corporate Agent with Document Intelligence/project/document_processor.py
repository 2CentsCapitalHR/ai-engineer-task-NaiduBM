#!/usr/bin/env python3
"""
ADGM Corporate Agent - Document Processing Backend
Handles .docx document analysis, compliance checking, and RAG integration
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
from datetime import datetime

# Document processing
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn

# AI and RAG components
import openai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.llms import OpenAI

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ComplianceIssue:
    document: str
    section: str
    issue: str
    severity: str  # 'High', 'Medium', 'Low'
    suggestion: str
    reference: str = ""

@dataclass
class DocumentAnalysis:
    process: str
    documents_uploaded: int
    required_documents: int
    missing_documents: List[str]
    issues_found: List[ComplianceIssue]
    confidence_score: float = 0.0

class ADGMComplianceRules:
    """ADGM compliance rules and validation logic"""
    
    INCORPORATION_DOCUMENTS = [
        "Articles of Association",
        "Memorandum of Association", 
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ]
    
    LICENSING_DOCUMENTS = [
        "FSRA License Application",
        "Business Plan",
        "Compliance Manual",
        "Risk Management Framework",
        "Key Personnel CVs"
    ]
    
    EMPLOYMENT_DOCUMENTS = [
        "Employment Contract",
        "Educational Certificates",
        "Experience Certificates",
        "Medical Certificate",
        "Passport Copy"
    ]
    
    COMPLIANCE_RULES = {
        "jurisdiction": {
            "pattern": r"ADGM|Abu Dhabi Global Market",
            "severity": "High",
            "message": "Document must specify ADGM jurisdiction",
            "reference": "ADGM Companies Regulations 2020, Article 6"
        },
        "ubo_declaration": {
            "pattern": r"Ultimate Beneficial Owner|UBO",
            "severity": "High", 
            "message": "UBO declaration required for AML compliance",
            "reference": "ADGM AML Rules 2019, Rule 3.2"
        },
        "registered_office": {
            "pattern": r"registered office.*ADGM",
            "severity": "High",
            "message": "Registered office must be in ADGM",
            "reference": "ADGM Companies Regulations 2020, Article 29"
        }
    }

class DocumentProcessor:
    """Main document processing and analysis engine"""
    
    def __init__(self, openai_api_key: str):
        self.openai_api_key = openai_api_key
        openai.api_key = openai_api_key
        
        # Initialize RAG components
        self.embeddings = OpenAIEmbeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        # Load ADGM knowledge base
        self.knowledge_base = self._initialize_knowledge_base()
        self.compliance_rules = ADGMComplianceRules()
        
    def _initialize_knowledge_base(self) -> Chroma:
        """Initialize RAG knowledge base with ADGM documents"""
        
        # ADGM reference texts (in production, load from actual documents)
        adgm_texts = [
            """
            ADGM Companies Regulations 2020 - Article 6: Jurisdiction
            All companies incorporated in ADGM are subject to ADGM jurisdiction.
            Legal disputes must be resolved through ADGM Courts unless otherwise specified.
            """,
            """
            ADGM AML Rules 2019 - Rule 3.2: Ultimate Beneficial Ownership
            All entities must maintain accurate records of Ultimate Beneficial Owners (UBO).
            UBO information must be declared and updated within specified timeframes.
            """,
            """
            ADGM Companies Regulations 2020 - Article 29: Registered Office
            Every company must maintain a registered office address within ADGM.
            The registered office must be accessible during business hours.
            """,
            """
            ADGM Companies Regulations 2020 - Article 155: Directors
            At least one director must be a natural person.
            Directors must meet fit and proper requirements as specified by ADGM.
            """
        ]
        
        # Split texts and create embeddings
        docs = self.text_splitter.create_documents(adgm_texts)
        
        # Create vector store
        vectorstore = Chroma.from_documents(
            documents=docs,
            embedding=self.embeddings,
            persist_directory="./adgm_knowledge_base"
        )
        
        return vectorstore
    
    def analyze_documents(self, file_paths: List[str]) -> DocumentAnalysis:
        """Analyze uploaded documents for ADGM compliance"""
        
        logger.info(f"Analyzing {len(file_paths)} documents")
        
        # Extract text from documents
        documents_content = {}
        for file_path in file_paths:
            try:
                doc = Document(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                documents_content[os.path.basename(file_path)] = content
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                continue
        
        # Identify process type
        process_type = self._identify_process(documents_content)
        
        # Check document completeness
        required_docs = self._get_required_documents(process_type)
        uploaded_docs = list(documents_content.keys())
        missing_docs = self._find_missing_documents(uploaded_docs, required_docs)
        
        # Perform compliance analysis
        issues = []
        for doc_name, content in documents_content.items():
            doc_issues = self._analyze_document_compliance(doc_name, content)
            issues.extend(doc_issues)
        
        # Create analysis result
        analysis = DocumentAnalysis(
            process=process_type,
            documents_uploaded=len(uploaded_docs),
            required_documents=len(required_docs),
            missing_documents=missing_docs,
            issues_found=issues,
            confidence_score=self._calculate_confidence_score(issues)
        )
        
        logger.info(f"Analysis complete: {len(issues)} issues found")
        return analysis
    
    def _identify_process(self, documents_content: Dict[str, str]) -> str:
        """Identify the legal process based on document content"""
        
        all_content = " ".join(documents_content.values()).lower()
        
        if any(term in all_content for term in ["incorporation", "articles of association", "memorandum"]):
            return "Company Incorporation"
        elif any(term in all_content for term in ["fsra", "financial services", "license"]):
            return "Financial Services Licensing"
        elif any(term in all_content for term in ["employment", "visa", "work permit"]):
            return "Employment Visa Application"
        else:
            return "General Compliance Review"
    
    def _get_required_documents(self, process_type: str) -> List[str]:
        """Get required documents for a specific process"""
        
        if process_type == "Company Incorporation":
            return self.compliance_rules.INCORPORATION_DOCUMENTS
        elif process_type == "Financial Services Licensing":
            return self.compliance_rules.LICENSING_DOCUMENTS
        elif process_type == "Employment Visa Application":
            return self.compliance_rules.EMPLOYMENT_DOCUMENTS
        else:
            return []
    
    def _find_missing_documents(self, uploaded: List[str], required: List[str]) -> List[str]:
        """Find missing required documents"""
        
        missing = []
        for req_doc in required:
            # Simple matching - in production, use more sophisticated matching
            if not any(req_doc.lower() in uploaded_doc.lower() for uploaded_doc in uploaded):
                missing.append(req_doc)
        
        return missing
    
    def _analyze_document_compliance(self, doc_name: str, content: str) -> List[ComplianceIssue]:
        """Analyze a single document for compliance issues"""
        
        issues = []
        
        # Rule-based compliance checking
        for rule_name, rule_config in self.compliance_rules.COMPLIANCE_RULES.items():
            import re
            if not re.search(rule_config["pattern"], content, re.IGNORECASE):
                issue = ComplianceIssue(
                    document=doc_name,
                    section="General",
                    issue=rule_config["message"],
                    severity=rule_config["severity"],
                    suggestion=f"Add {rule_name} clause as per {rule_config['reference']}",
                    reference=rule_config["reference"]
                )
                issues.append(issue)
        
        # AI-powered analysis using RAG
        ai_issues = self._ai_compliance_analysis(doc_name, content)
        issues.extend(ai_issues)
        
        return issues
    
    def _ai_compliance_analysis(self, doc_name: str, content: str) -> List[ComplianceIssue]:
        """Use AI and RAG for advanced compliance analysis"""
        
        try:
            # Create QA chain with knowledge base
            qa_chain = RetrievalQA.from_chain_type(
                llm=OpenAI(temperature=0),
                chain_type="stuff",
                retriever=self.knowledge_base.as_retriever()
            )
            
            # Query for compliance issues
            query = f"""
            Analyze the following document content for ADGM compliance issues:
            Document: {doc_name}
            Content: {content[:2000]}...
            
            Identify any legal issues, missing clauses, or non-compliance with ADGM regulations.
            """
            
            result = qa_chain.run(query)
            
            # Parse AI response (simplified - in production, use structured output)
            if "issue" in result.lower():
                issue = ComplianceIssue(
                    document=doc_name,
                    section="AI Analysis",
                    issue="Potential compliance issue identified",
                    severity="Medium",
                    suggestion=result[:200] + "...",
                    reference="AI Analysis with ADGM Knowledge Base"
                )
                return [issue]
            
        except Exception as e:
            logger.error(f"AI analysis error: {e}")
        
        return []
    
    def _calculate_confidence_score(self, issues: List[ComplianceIssue]) -> float:
        """Calculate confidence score based on analysis results"""
        
        if not issues:
            return 0.95
        
        high_severity_count = sum(1 for issue in issues if issue.severity == "High")
        medium_severity_count = sum(1 for issue in issues if issue.severity == "Medium")
        
        # Simple scoring algorithm
        score = 1.0 - (high_severity_count * 0.2 + medium_severity_count * 0.1)
        return max(0.0, min(1.0, score))
    
    def add_comments_to_document(self, file_path: str, issues: List[ComplianceIssue]) -> str:
        """Add compliance comments to the original document"""
        
        try:
            doc = Document(file_path)
            
            # Add comments for each issue (simplified implementation)
            for issue in issues:
                # Find relevant paragraph and add comment
                for paragraph in doc.paragraphs:
                    if len(paragraph.text) > 50:  # Add to substantial paragraphs
                        # Add comment (this is a simplified approach)
                        comment_text = f"COMPLIANCE NOTE: {issue.issue} - {issue.suggestion}"
                        
                        # Create a new paragraph with the comment
                        comment_paragraph = doc.add_paragraph()
                        comment_run = comment_paragraph.add_run(f"[COMMENT: {comment_text}]")
                        comment_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                        break
            
            # Save modified document
            output_path = file_path.replace('.docx', '_reviewed.docx')
            doc.save(output_path)
            
            logger.info(f"Comments added to document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error adding comments: {e}")
            return file_path

def main():
    """Main function for testing"""
    
    # Initialize processor
    processor = DocumentProcessor(openai_api_key="your-openai-api-key")
    
    # Example usage
    test_files = ["sample_articles.docx", "sample_memorandum.docx"]
    
    # Analyze documents
    analysis = processor.analyze_documents(test_files)
    
    # Print results
    print(json.dumps(asdict(analysis), indent=2, default=str))

if __name__ == "__main__":
    main()